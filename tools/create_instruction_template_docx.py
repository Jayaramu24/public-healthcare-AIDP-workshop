from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "codex_workshop_instruction_template.md"
OUTPUT = ROOT / "codex_workshop_instruction_template.docx"

DEEP = RGBColor(0x14, 0x37, 0x35)
MUTED = RGBColor(0x56, 0x64, 0x60)
ORACLE = RGBColor(0xC8, 0x32, 0x32)
SOFT = "F4F8F6"
LINE = "D9E4DF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        el = tc_mar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = DEEP
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11.5)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(18)

    for style_name, size in [("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11)]:
        st = doc.styles[style_name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = DEEP if style_name != "Heading 3" else ORACLE
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)

    if "CodeBlock" not in doc.styles:
        code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(8.8)
        code_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.18)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(8)
        code_style.paragraph_format.line_spacing = 1.08


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Codex Workshop Instruction Template")

    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run(
        "Reusable instruction pack for building AIDP, Autonomous AI Lakehouse, "
        "OAC, ML, and AI agent workshop assets across domains."
    )

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.columns[0].width = Inches(2.5)
    meta.columns[1].width = Inches(3.4)

    values = [
        ("Designed for", "Finance, Banking, Insurance, Manufacturing, Retail, Public Sector"),
        ("Output types", "HTML page, Markdown guide, PDF guide, notebooks, datasets, optional lab assets"),
    ]
    for r, (label, value) in enumerate(values):
        for c, text in enumerate((label, value)):
            cell = meta.cell(r, c)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if c == 0:
                set_cell_shading(cell, SOFT)
                for para in cell.paragraphs:
                    if para.runs:
                        para.runs[0].bold = True
                        para.runs[0].font.color.rgb = DEEP

    doc.add_paragraph(
        "Use this document as the reusable request template when you want Codex "
        "to generate a complete workshop package similar to the Public Healthcare "
        "AIDP Workshop, but for another domain.",
    )

    doc.add_page_break()


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = False
        p.add_run(line)
        if line == lines[-1]:
            p.paragraph_format.space_after = Pt(10)


def add_section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section = doc.sections[-1]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("How to Use This Template", level=1)
    for item in [
        "Fill in the input sections with domain-specific details.",
        "Keep constraints explicit, especially around data volume, schema style, and workshop flow.",
        "Provide screenshots, links, or reference assets whenever look and feel matters.",
        "Paste the completed input section into the master prompt later in this document.",
        "Ask Codex to propose first if you want to review the structure before it starts updating assets.",
    ]:
        add_number(doc, item)

    doc.add_heading("Part 1 - Core Inputs to Provide", level=1)

    sections = [
        (
            "A. Workshop Identity",
            [
                "exact workshop title",
                "domain name",
                "target audience",
                "workshop duration",
                "delivery mode: instructor-led, self-paced, or hybrid",
                "whether the output is customer-facing, internal execution-facing, or both",
            ],
            [
                "Workshop title: `Corporate Banking AIDP Workshop`",
                "Audience: `solution engineers, data architects, analytics leads`",
                "Delivery mode: `hybrid`",
            ],
        ),
        (
            "B. Business Use-Case Scope",
            [
                "the primary business scenario",
                "3 to 6 sub-use-cases to reflect in the data",
                "the main business questions the workshop should answer",
                "whether there should be one guided use-case or multiple use-case lanes",
                "whether one use-case should be instructor-led and another DIY",
            ],
            [
                "finance: close process, spend controls, planning variance, procurement, treasury",
                "banking: onboarding, deposits, lending, collections, fraud, branch operations",
                "manufacturing: production, quality, maintenance, supplier risk, inventory, plant throughput",
            ],
        ),
        (
            "C. Technology Scope",
            [
                "use Oracle AI Data Platform: yes or no",
                "use Autonomous AI Lakehouse: yes or no",
                "use Oracle Analytics Cloud: yes or no",
                "include ML lab: yes or no",
                "include AI agent lab: yes or no",
                "include real-time or streaming lab: yes or no",
                "services to explicitly exclude",
            ],
            [],
        ),
        (
            "D. Data Inputs to Synthesize",
            [
                "maximum number of CSV datasets",
                "whether to include one JSON dataset",
                "whether to include one spatial dataset",
                "whether to include one document dataset",
                "minimum document size, for example 10 or more pages",
                "date range for the synthetic data",
                "operating footprint, for example regions, branches, plants, facilities, or business units",
            ],
            [
                "Recommended compact pattern:",
                "max 5 CSV datasets",
                "1 JSON dataset",
                "1 spatial dataset",
                "1 document dataset",
            ],
        ),
        (
            "E. Data Modeling Expectations",
            [
                "what Bronze should contain",
                "what Silver should contain",
                "what Gold should contain",
                "whether Gold must be a star schema or snowflake schema",
                "whether Gold should contain one fact table with dimensions or multiple stars",
                "whether one Gold star should be instructor-led and another should be DIY",
            ],
            [
                "Bronze should preserve raw structure and add ingestion metadata only",
                "Silver should type, validate, conform, and derive operational fields",
                "Gold should publish business-serving star schema objects for OAC",
            ],
        ),
        (
            "F. Dashboard Requirements",
            [
                "number of dashboards or canvases",
                "required dashboard topics",
                "whether dashboard titles should appear as horizontal tabs",
                "whether working filters are required",
                "whether spatial insight should use a map",
                "whether an AI assistant side panel pattern is needed",
                "screenshot or wireframe references for desired OAC look and feel",
            ],
            [],
        ),
        (
            "G. Lab Design Requirements",
            [
                "which setup steps are `Admin step`",
                "which hands-on steps are `Participant step`",
                "whether setup should be folded into lab execution steps",
                "whether Codex should provide OCI UI step-by-step guidance",
                "whether screenshots are required for AIDP, AI Lakehouse, OAC, ML, and Agents",
            ],
            [
                "Recommended rule:",
                "do not keep setup as a disconnected preamble",
                "fold setup into the labs",
                "label each step as `Admin step` or `Participant step`",
            ],
        ),
        (
            "H. Output Assets Required",
            [
                "HTML workshop page",
                "Markdown workshop guide",
                "downloadable PDF guide",
                "medallion architecture diagrams",
                "Bronze to Silver transformation diagrams",
                "Silver to Gold flow diagrams",
                "Gold star schema diagrams",
                "Bronze notebook",
                "Silver notebook",
                "Gold notebook",
                "ML notebook or `.py` file",
                "optional lab markdowns",
                "screenshots or mock visuals",
            ],
            [],
        ),
        (
            "I. Style and Presentation Preferences",
            [
                "preferred tone: professional, workshop-demo, customer-facing, internal execution",
                "whether to align wording to official Oracle documentation",
                "whether generic workshop names should be avoided",
                "whether diagrams should be executive-style, technical, or hybrid",
                "whether visuals should use swimlanes, block diagrams, RDBMS-style layouts, or star-schema diagrams",
            ],
            [],
        ),
        (
            "J. Reference Material to Align To",
            [
                "Oracle documentation links",
                "Oracle blog links",
                "LiveLabs links",
                "screenshots or mockups",
                "existing HTML pages",
                "PowerPoint or deck references",
                "any prior workshop assets that should be mirrored",
            ],
            [],
        ),
    ]

    for title, bullets, extras in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph("Provide:")
        for item in bullets:
            add_bullet(doc, item)
        for extra in extras:
            if extra.endswith(":"):
                doc.add_paragraph(extra)
            else:
                add_bullet(doc, extra)

    add_section_break(doc)

    doc.add_heading("Part 2 - Constraints You Should State Explicitly", level=1)
    for item in [
        "use synthetic data only",
        "keep raw datasets limited and intentional",
        "use medallion architecture in AIDP",
        "define Bronze to Silver logic clearly",
        "define Silver to Gold logic clearly",
        "use a Gold star schema unless another schema style is explicitly requested",
        "use OAC on the instructor-led Gold flow unless otherwise specified",
        "use map-based spatial insight if spatial data is included",
        "use the document dataset for vectorization and grounded chat if a document is included",
        "add ML and AI agent labs only if requested",
        "label all setup steps as `Admin step` or `Participant step`",
        "fold environment setup into the labs",
        "generate HTML and PDF outputs if the workshop is intended for sharing",
        "use separate notebooks for Bronze, Silver, Gold, and ML when relevant",
        "align AIDP, AI Lakehouse, OAC, ML, and Agent guidance to the provided Oracle docs and blogs",
        "when major structural changes are requested, propose first if the user asks for review before updating",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Part 3 - Fill-In Template", level=1)
    doc.add_paragraph("Copy the block below and fill in the values.")
    fill_template = SOURCE.read_text(encoding="utf-8").split("## Part 3 - Fill-In Template", 1)[1]
    fill_template = fill_template.split("## Part 4 - Master Prompt to Give Codex", 1)[0]
    code = fill_template.split("```text", 1)[1].split("```", 1)[0].strip("\n")
    add_code_block(doc, code.splitlines())

    doc.add_heading("Part 4 - Master Prompt to Give Codex", level=1)
    doc.add_paragraph("Once you fill in the template, use the prompt below.")
    master = SOURCE.read_text(encoding="utf-8").split("## Part 4 - Master Prompt to Give Codex", 1)[1]
    master = master.split("## Part 5 - Recommended Minimum Input Set", 1)[0]
    master_code = master.split("```text", 1)[1].split("```", 1)[0].strip("\n")
    add_code_block(doc, master_code.splitlines())

    doc.add_heading("Part 5 - Recommended Minimum Input Set", level=1)
    for item in [
        "exact workshop title",
        "domain and target audience",
        "primary scenario plus 3 to 6 sub-use-cases",
        "raw data constraints",
        "Gold schema choice",
        "instructor-led and DIY split, if any",
        "dashboard reference screenshot",
        "which labs are admin versus participant",
        "Oracle docs and blogs to align to",
        "required output assets",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Part 6 - Optional Review-First Prompt", level=1)
    review = SOURCE.read_text(encoding="utf-8").split("## Part 6 - Optional Review-First Prompt", 1)[1]
    review = review.split("## Part 7 - Notes for Best Results", 1)[0]
    review_code = review.split("```text", 1)[1].split("```", 1)[0].strip("\n")
    add_code_block(doc, review_code.splitlines())

    doc.add_heading("Part 7 - Notes for Best Results", level=1)
    for item in [
        "Provide exact names when naming matters.",
        "Attach screenshots whenever look and feel matters.",
        "Constrain the raw data early to avoid unnecessary dataset sprawl.",
        "Decide early whether the workshop is primarily technical, business-facing, or demo-oriented.",
        "Decide early whether the agent experience should use official OCI flows, representative workshop visuals, or both.",
        "For customer workshops, ask Codex to keep diagrams intuitive and presentation-ready, not only technically correct.",
    ]:
        add_bullet(doc, item)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if footer.runs:
            run = footer.runs[0]
        else:
            run = footer.add_run()
        run.text = "Codex Workshop Instruction Template"
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
